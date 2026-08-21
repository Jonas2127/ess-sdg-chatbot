"""
Word Exporter for ESS RAG Chatbot Conversations
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os


class WordExporter:
    """Export conversations to Word format"""
    
    def __init__(self):
        # Use absolute path to ensure logo is found regardless of working directory
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        self.logo_path = os.path.join(base_dir, "assets", "ess_logo_fixed.png")
    
    def export_conversation(self, messages, filename=None):
        """
        Export conversation to Word document
        
        Args:
            messages: List of message dictionaries with 'role', 'content', 'metadata'
            filename: Output filename (optional, auto-generated if not provided)
            
        Returns:
            dict: {'success': bool, 'filename': str, 'error': str}
        """
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"ess_conversation_{timestamp}.docx"
            
            # Ensure output directory exists
            output_dir = "exports"
            os.makedirs(output_dir, exist_ok=True)
            filepath = os.path.join(output_dir, filename)
            
            # Create Word document
            doc = Document()
            
            # Set default font to Times New Roman
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            
            # Add ESS logo at the top (centered)
            if os.path.exists(self.logo_path):
                try:
                    logo_para = doc.add_paragraph()
                    logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    logo_run = logo_para.add_run()
                    logo_run.add_picture(self.logo_path, width=Inches(1.5))
                    doc.add_paragraph()  # Spacing after logo
                except Exception as logo_err:
                    # Log error but continue without logo
                    print(f"Warning: Could not add logo to Word document: {logo_err}")
            else:
                print(f"Warning: Logo file not found at: {self.logo_path}")
            
            # Add header
            title = doc.add_heading('Ethiopian Statistics Service', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title.runs[0]
            title_run.font.name = 'Times New Roman'
            title_run.font.color.rgb = RGBColor(74, 222, 128)  # Green
            
            subtitle = doc.add_paragraph('RAG Chatbot Conversation Export')
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.name = 'Times New Roman'
            subtitle_run.font.size = Pt(12)
            subtitle_run.font.color.rgb = RGBColor(107, 155, 209)  # Blue
            
            date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            date_run = date_para.runs[0]
            date_run.font.name = 'Times New Roman'
            date_run.font.size = Pt(10)
            date_run.font.color.rgb = RGBColor(148, 163, 184)  # Gray
            
            doc.add_paragraph()  # Spacing
            doc.add_paragraph('_' * 80)  # Separator
            doc.add_paragraph()
            
            # Add Q&A pairs
            for idx, msg in enumerate(messages, 1):
                if msg['role'] == 'user':
                    # Question
                    q_para = doc.add_paragraph()
                    q_run = q_para.add_run(f"Q{idx}: {msg['content']}")
                    q_run.font.name = 'Times New Roman'
                    q_run.font.bold = True
                    q_run.font.size = Pt(12)
                    q_run.font.color.rgb = RGBColor(74, 222, 128)  # Green
                    
                elif msg['role'] == 'assistant':
                    # Answer
                    a_para = doc.add_paragraph()
                    a_label = a_para.add_run("A: ")
                    a_label.font.name = 'Times New Roman'
                    a_label.font.bold = True
                    a_label.font.size = Pt(12)
                    
                    a_content = a_para.add_run(msg['content'])
                    a_content.font.name = 'Times New Roman'
                    a_content.font.size = Pt(12)
                    a_content.font.color.rgb = RGBColor(0, 0, 0)  # Black for better readability
                    
                    # Metadata
                    if 'metadata' in msg:
                        metadata = msg['metadata']
                        meta_para = doc.add_paragraph()
                        meta_text = f"⏱️ Response Time: {metadata.get('time', 0):.2f}s  |  "
                        meta_text += f"🔧 Engines: {', '.join(metadata.get('engines', []))}  |  "
                        meta_text += f"📚 Sources: {metadata.get('sources', 0)}"
                        
                        meta_run = meta_para.add_run(meta_text)
                        meta_run.font.name = 'Times New Roman'
                        meta_run.font.size = Pt(9)
                        meta_run.font.italic = True
                        meta_run.font.color.rgb = RGBColor(148, 163, 184)  # Gray
                    
                    # Add separator
                    doc.add_paragraph()
                    doc.add_paragraph('─' * 80)
                    doc.add_paragraph()
            
            # Add footer
            doc.add_paragraph()
            footer = doc.add_paragraph()
            footer_run = footer.add_run(
                'This document contains AI-generated responses based on '
                'Ethiopian Statistics Service data and reports.'
            )
            footer_run.font.name = 'Times New Roman'
            footer_run.font.size = Pt(9)
            footer_run.font.italic = True
            footer_run.font.color.rgb = RGBColor(148, 163, 184)
            
            # Save document
            doc.save(filepath)
            
            return {
                'success': True,
                'filename': filename,
                'filepath': filepath,
                'error': None
            }
            
        except Exception as e:
            return {
                'success': False,
                'filename': None,
                'filepath': None,
                'error': str(e)
            }
